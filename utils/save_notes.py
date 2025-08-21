from docx import Document

def save_as_docx(summary_text, filepath):
    doc = Document()
    doc.add_heading("Video Notes", 0)
    doc.add_paragraph(summary_text)
    doc.save(filepath)
